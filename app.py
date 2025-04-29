import streamlit as st
st.set_page_config(layout="wide", page_title="Vojenský obor FTVS UK")

import os
import datetime
import pandas as pd
from io import BytesIO
from docx.shared import Pt
from docx import Document
from docxcompose.composer import Composer
from streamlit_quill import st_quill

from supabase import create_client, Client
SUPABASE_URL = st.secrets["supabase"]["supabase_url"]
SUPABASE_KEY = st.secrets["supabase"]["supabase_key"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Pro ruční rerun
from streamlit.runtime.scriptrunner import RerunException, RerunData

# Dummy moduly – pokud nejsou k dispozici
try:
    import dpp
except ModuleNotFoundError:
    class dpp:
        @staticmethod
        def run_dpp():
            st.info("Modul DPP není k dispozici.")
try:
    import pri
except ModuleNotFoundError:
    class pri:
        @staticmethod
        def run_pri(year):
            st.info("Modul PRI není k dispozici.")
try:
    import zsc
except ModuleNotFoundError:
    class zsc:
        @staticmethod
        def run_zsc():
            st.info("Modul ZSC není k dispozici.")

# Sidebar – nastavení položek vyhodnocení
items = [
    "Souhrnný přehled APVVP","VŠ vzdělávání","Přijímačky","Akreditace",
    "Vědecká činnost","Zahraniční spolupráce","Personální oblast",
    "Logistika","Ekonomika","Odborné kurzy","Vojenská příprava","Jazykové vzdělávání"
]
with st.sidebar.expander("Nastavení položek vyhodnocení"):
    for itm in items:
        st.markdown(f"#### {itm}")
        st.checkbox("Zobrazit", key=f"include_{itm}")
        st.checkbox("Hotovo", key=f"finished_{itm}")
    st.markdown("---")
    st.checkbox("Zobrazit celkové vyhodnocení", key="include_celkovy")

# Autentizace
PASSWORD = st.secrets["app"]["login_password"]
if not st.session_state.get("authenticated", False):
    st.title("Přihlášení")
    pwd = st.text_input("Zadejte heslo", type="password")
    if st.button("Přihlásit se"):
        if pwd == PASSWORD:
            st.session_state.authenticated = True
            st.success("Přihlášení proběhlo úspěšně!")
        else:
            st.error("Špatné heslo!")
    st.stop()

# Hlavička
col1, col2 = st.columns([1, 6])
with col1:
    st.image("Logo.png", width=50)
with col2:
    st.markdown("<h1 style='margin-bottom:0;'>Vojenský obor FTVS UK</h1>", unsafe_allow_html=True)

# Custom CSS
st.markdown("""
<style>
.ql-editor { font-family:'Times New Roman', serif; font-size:14px; }
.stTextInput>div>div>input { max-width:300px; }
</style>
""", unsafe_allow_html=True)

def format_table(doc_table, font_size=10):
    doc_table.style = "Table Grid"
    for row in doc_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)

# Funkce pro souhrn studentů
def run_summary():
    st.header("Souhrn všech studentů")
    try:
        resp = supabase.table("students").select("*").execute()
        students = resp.data or []
    except Exception as e:
        st.error(f"Chyba při načítání studentů: {e}")
        return
    if not students:
        st.info("Žádní studenti nejsou zaregistrováni.")
        return

    base_cols = ["Hodnost","Jméno","Příjmení","Ročník","Kohorta"]
    desired = [
        "TaD-I","TaD-II","TaD-III","TaD-1","TaD-2","TaD-3",
        "ZSTP-I","ZSTP-II","ZSTP-III","ZSTP-1","ZSTP-2",
        "STP-I","STP-II","STP-III","STP-1","STP-2",
        "Kurz BZ-I","Kurz BZ-II","Kurz BZ-III","Kurz BZ-IV",
        "Kurz VL-I","Kurz VL-II","Kurz VL-III","Kurz VL-IV",
        "Kurz PSL","Kurz PSL-I","Kurz PSL-II",
        "Kurz ZP-I","Kurz ZP-II","Kurz VPL-I","Kurz VPL-II",
        "Kurz STP-I","Kurz STP-II",
        "Instr. BZ","Instr. VL","Instr. PSL","Instr. ZP","Instr. VPL"
    ]
    mapping = {
        "Teorie a didaktika AČR-I":"TaD-I","Teorie a didaktika AČR-II":"TaD-II","Teorie a didaktika AČR-III":"TaD-III",
        "Teorie a didaktika AČR-1":"TaD-1","Teorie a didaktika AČR-2":"TaD-2","Teorie a didaktika AČR-3":"TaD-3",
        "Základy STP-I":"ZSTP-I","Základy STP-II":"ZSTP-II","Základy STP-III":"ZSTP-III",
        "Základy STP-1":"ZSTP-1","Základy STP-2":"ZSTP-2",
        "Speciální TP-I":"STP-I","Speciální TP-II":"STP-II","Speciální TP-III":"STP-III",
        "Speciální TP-1":"STP-1","Speciální TP-2":"STP-2",
        "Kurz BZ-I":"Kurz BZ-I","Kurz BZ-II":"Kurz BZ-II","Kurz BZ-III":"Kurz BZ-III","Kurz BZ-IV":"Kurz BZ-IV",
        "Kurz VL-I":"Kurz VL-I","Kurz VL-II":"Kurz VL-II","Kurz VL-III":"Kurz VL-III","Kurz VL-IV":"Kurz VL-IV",
        "Kurz PSL":"Kurz PSL","Kurz PSL-I":"Kurz PSL-I","Kurz PSL-II":"Kurz PSL-II",
        "Kurz ZP-I":"Kurz ZP-I","Kurz ZP-II":"Kurz ZP-II",
        "Kurz VPL-I":"Kurz VPL-I","Kurz VPL-II":"Kurz VPL-II",
        "Kurz STP-I":"Kurz STP-I","Kurz STP-II":"Kurz STP-II",
        "BZ-IV Instruktor":"Instr. BZ","VL-IV Instruktor":"Instr. VL",
        "PSL-II Instruktor":"Instr. PSL","ZP-II Instruktor":"Instr. ZP","VPL-II Instruktor":"Instr. VPL"
    }

    def is_done(detail):
        if not isinstance(detail, dict):
            return False
        if "instruktor" in detail:
            return bool(detail.get("instruktor"))
        if "completed" in detail:
            return bool(detail.get("completed"))
        if detail.get("grade","").strip() != "":
            return True
        for v in detail.values():
            if is_done(v):
                return True
        return False

    def extract(s):
        row = {col: "NE" for col in desired}
        subj = s.get("subjects", {})
        for sem in ("zimni", "letni"):
            for full, det in subj.get(sem, {}).items():
                ab = mapping.get(full)
                if ab:
                    row[ab] = "ANO" if is_done(det) else "NE"
        return row

    records = []
    for s in students:
        base = {
            "Hodnost":  s.get("hodnost", ""),
            "Jméno":    s.get("first_name", ""),
            "Příjmení": s.get("last_name", ""),
            "Ročník":   s.get("cohort", ""),
            "Kohorta":  s.get("study_type", "")
        }
        records.append({**base, **extract(s)})

    df = pd.DataFrame(records, columns=base_cols + desired)
    # Seřazení podle ročníků
    order = {"1. Bc.":0, "2. Bc.":1, "3. Bc.":2, "1. Mgr.":3, "2. Mgr.":4}
    df["__order"] = df["Ročník"].map(order).fillna(len(order))
    df = df.sort_values(["__order", "Příjmení", "Jméno"])
    df = df.drop(columns="__order")

    st.dataframe(df, use_container_width=True)
    if st.button("Exportovat do Excelu"):
        buf = BytesIO()
        with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Studenti")
        st.download_button(
            "Stáhnout Excel",
            buf.getvalue(),
            "Studenti_souhrn.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# Evaluace
def load_evaluations():
    r = supabase.table("evaluations").select("data").eq("id", 1).execute()
    return r.data[0]["data"] if r.data and r.data[0].get("data") else {}

def save_evaluations(data):
    return supabase.table("evaluations").update({"data": data}).eq("id", 1).execute()

if "evaluations" not in st.session_state:
    st.session_state.evaluations = load_evaluations()

evaluation_periods = {
    "1. čtvrtletí": ("1. leden","31. březen"),
    "1. pololetí":  ("1. leden","30. červen"),
    "3. čtvrtletí":("1. červenec","31. září"),
    "Celý rok":     ("1. leden","31. prosinec")
}

# Hlavní záložky
tabs = st.tabs([
    "Vyhodnocení VO FTVS UK",
    "Historie vyhodnocení",
    "DPP",
    "PR-I",
    "ZSC",
    "Student"
])

with tabs[0]:
    st.header("Vyhodnocení VO FTVS UK")
    year = st.number_input("Rok",2000,2100,datetime.datetime.now().year)
    period = st.selectbox("Vyberte období", list(evaluation_periods.keys()))
    pr = evaluation_periods[period]
    st.write(f"{period} ({pr[0]} až {pr[1]})")
    key = f"{year}_{period}"
    saved = st.session_state.evaluations.get(key, {})
    to_eval = items if st.session_state.include_celkovy else [
        i for i in items if st.session_state.get(f"include_{i}")
    ]
    if not to_eval:
        st.info("Vyberte položky k vyhodnocení.")
    else:
        st.markdown("### Vyplňte vyhodnocení")
        new = {}
        for it in to_eval:
            st.markdown(f"### {it}")
            if it == "Souhrnný přehled APVVP":
                st.markdown("##### Nahrajte excel soubor s přehledem APVVP:")
                saved_table = saved.get(it, {}).get("table")
                if saved_table:
                    st.info("Uložená tabulka:")
                    st.table(saved_table)
                uploaded_file = st.file_uploader(
                    "Vyberte soubor", type=["xlsx"], key=f"upload_apvvp_{year}_{period}"
                )
                table_data = None
                if uploaded_file is not None:
                    try:
                        df_tmp = pd.read_excel(uploaded_file)
                        df_tmp = df_tmp.fillna("").astype(str)
                        st.dataframe(df_tmp)
                        table_data = [list(df_tmp.columns)] + df_tmp.values.tolist()
                    except Exception:
                        st.error("Chyba při načítání excel souboru.")
                new[it] = {"table": table_data if table_data is not None else saved.get(it, {}).get("table")}
            elif it == "Ekonomika":
                st.markdown("##### Nahrajte excel soubor pro Ekonomiku:")
                saved_table = saved.get(it, {}).get("table")
                if saved_table:
                    st.info("Uložená tabulka:")
                    st.table(saved_table)
                uploaded_file = st.file_uploader(
                    "Vyberte soubor", type=["xlsx"], key=f"upload_ekonomika_{year}_{period}"
                )
                table_data = None
                if uploaded_file is not None:
                    try:
                        df_tmp = pd.read_excel(uploaded_file)
                        df_tmp = df_tmp.fillna("").astype(str)
                        st.dataframe(df_tmp)
                        table_data = [list(df_tmp.columns)] + df_tmp.values.tolist()
                    except Exception:
                        st.error("Chyba při načítání excel souboru pro Ekonomiku.")
                txt = st_quill(
                    key=f"eval_{year}_{period}_{it}", value=saved.get(it, {}).get("text", "")
                )
                fin = st.checkbox(
                    "Hotovo", key=f"fin_{year}_{period}_{it}", value=saved.get(it, {}).get("finished", False)
                )
                new[it] = {
                    "table": table_data if table_data is not None else saved.get(it, {}).get("table"),
                    "text": txt,
                    "finished": fin
                }
            else:
                txt = st_quill(
                    key=f"eval_{year}_{period}_{it}", value=saved.get(it, {}).get("text", "")
                )
                fin = st.checkbox(
                    "Hotovo", key=f"fin_{year}_{period}_{it}", value=saved.get(it, {}).get("finished", False)
                )
                new[it] = {"text": txt, "finished": fin}
        if st.button("Uložit vyhodnocení"):
            st.session_state.evaluations[key] = new
            save_evaluations(st.session_state.evaluations)
            st.success("Vyhodnocení uloženo!")
        st.subheader("Generování dokumentů")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generovat Word dokument", key="gen_word"):
                try:
                    intro_path = "upload/Úvodní strana vyhodnocení.docx"
                    if os.path.exists(intro_path):
                        intro_doc = Document(intro_path)
                    else:
                        st.error("Soubor Úvodní strana vyhodnocení.docx nebyl nalezen v adresáři upload.")
                        intro_doc = Document()
                    eval_doc = Document()
                    eval_doc.add_heading(f"Vyhodnocení za: {period} roku {year}", level=0)
                    eval_doc.add_page_break()
                    doc_items = items if st.session_state.get("include_celkovy", False) else to_eval
                    for item in doc_items:
                        # Use current in-memory evaluation data (including uploaded tables)
                        stored = st.session_state.evaluations.get(key, {})
                        data = new.get(item, stored.get(item, {}))
                        eval_doc.add_heading(item, level=2)
                        if data.get("table"):
                            table_data = data["table"]
                            rows, cols = len(table_data), len(table_data[0])
                            table = eval_doc.add_table(rows=rows, cols=cols)
                            for i, row in enumerate(table_data):
                                for j, cell in enumerate(row):
                                    table.cell(i, j).text = str(cell)
                            format_table(table)
                        if data.get("text"):
                            eval_doc.add_paragraph(data["text"])
                    composer = Composer(intro_doc)
                    composer.append(eval_doc)
                    buffer = BytesIO()
                    composer.save(buffer)
                    buffer.seek(0)
                    st.download_button("Stáhnout Word dokument", buffer.getvalue(), "Vyhodnoceni.docx")
                except Exception as e:
                    st.error(f"Chyba při generování dokumentu: {e}")

with tabs[1]:
    st.header("Historie vyhodnocení")
    hy = st.number_input("Rok historie",2000,2100,datetime.datetime.now().year,key="hist_year")
    hp = st.selectbox("Období historie", list(evaluation_periods.keys()),key="hist_period")
    hk = f"{hy}_{hp}"
    if hk in st.session_state.evaluations:
        st.subheader(f"Vyhodnocení za {hk}")
        for it,d in st.session_state.evaluations[hk].items():
            mark = " (hotovo)" if d.get("finished") else ""
            st.markdown(f"### {it}{mark}")
            st.write(d.get("text",""))
    else:
        st.info("Žádná data.")

with tabs[2]:
    dpp.run_dpp()
with tabs[3]:
    pri.run_pri(st.number_input("Rok PR‑I",2000,2100,datetime.datetime.now().year))
with tabs[4]:
    zsc.run_zsc()
with tabs[5]:
    cols = st.columns([8,1])
    cols[0].header("Student")
    if cols[1].button("Aktualizace"):
        raise RerunException(RerunData(st.query_params))

    student_tabs = st.tabs([
        "Vojenské předměty","Přidat studenta","Editace studenta","Souhrn","Absolventi"
    ])
    import student_1Bc, student_2Bc, student_3Bc, student_1Mgr, student_2Mgr, student
    with student_tabs[0]:
        cohort_tabs = st.tabs([
            "První ročník","Druhý ročník","Třetí ročník","Čtvrtý ročník","Pátý ročník"
        ])
        with cohort_tabs[0]:
            student_1Bc.run_student()
        with cohort_tabs[1]:
            student_2Bc.run_student()
        with cohort_tabs[2]:
            student_3Bc.run_student()
        with cohort_tabs[3]:
            student_1Mgr.run_student()
        with cohort_tabs[4]:
            student_2Mgr.run_student()
    with student_tabs[1]:
        student.run_add_student()
    with student_tabs[2]:
        student.run_edit_student()
    with student_tabs[3]:
        run_summary()
    with student_tabs[4]:
        student.run_graduates()
